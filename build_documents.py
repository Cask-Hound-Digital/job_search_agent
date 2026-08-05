import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

# Generic Document Builder for Tailored Resumes and Cover Letters
def make_docx_resume(profile, job_info, filepath):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(profile["name"].upper())
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p_c = doc.add_paragraph()
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c.add_run(f"{profile['location']} | {profile['phone']} | {profile['email']} | {profile['linkedin']}")

    doc.save(filepath)
    print(f"Generated DOCX Resume: {filepath}")

if __name__ == '__main__':
    print("Run build_documents.py with your configured candidate profile.")
